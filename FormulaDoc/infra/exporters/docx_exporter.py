from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from core.document_models import (
    DocumentData,
    EquationBlock,
    EquationRun,
    HeadingBlock,
    ImageBlock,
    PageBreakBlock,
    ParagraphBlock,
    TableBlock,
    TextRun,
)
from core.document_parsers import document_from_markdown
from infra.equation.omml_converter import OmmlConverter


def export_markdown_to_docx(markdown: str, output_path: Path) -> None:
    export_document_to_docx(document_from_markdown(markdown), output_path)


def export_document_to_docx(document_data: DocumentData, output_path: Path) -> None:
    document = Document()
    configure_document(document)
    converter = OmmlConverter()

    for block in document_data.blocks:
        if isinstance(block, HeadingBlock):
            add_heading_block(document, block)
        elif isinstance(block, ParagraphBlock):
            add_paragraph_block(document, converter, block)
        elif isinstance(block, EquationBlock):
            add_equation_block(document, converter, block)
        elif isinstance(block, TableBlock):
            add_table_block(document, block)
        elif isinstance(block, ImageBlock):
            add_image_block(document, block)
        elif isinstance(block, PageBreakBlock):
            document.add_page_break()

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "宋体"
    normal_style.font.size = Pt(12)
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def set_run_font(run, font_name: str, font_size: int, bold: bool = False) -> None:
    run.bold = bold
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def format_body_paragraph(paragraph) -> None:
    paragraph.paragraph_format.first_line_indent = Pt(24)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_heading(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(10)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    set_run_font(run, font_name="黑体", font_size=16)


def add_heading_block(document: Document, block: HeadingBlock) -> None:
    if block.level <= 1:
        add_heading(document, block.text)
        return
    paragraph = document.add_heading(level=min(max(block.level - 1, 0), 9))
    run = paragraph.add_run(block.text)
    set_run_font(run, font_name="黑体", font_size=max(12, 18 - block.level))


def add_text_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    format_body_paragraph(paragraph)
    run = paragraph.add_run(text)
    set_run_font(run, font_name="宋体", font_size=12)


def add_paragraph_block(document: Document, converter: OmmlConverter, block: ParagraphBlock) -> None:
    paragraph = document.add_paragraph()
    format_body_paragraph(paragraph)
    for part in block.parts:
        if isinstance(part, TextRun):
            run = paragraph.add_run(part.text)
            set_run_font(run, font_name="宋体", font_size=12)
        elif isinstance(part, EquationRun):
            run = paragraph.add_run()
            run._element.append(converter.to_omml(part.latex))


def add_equation_block(document: Document, converter: OmmlConverter, block: EquationBlock) -> None:
    if not block.latex.strip():
        return
    if block.number:
        add_numbered_equation(document, converter, block.latex, block.number)
        return

    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run._element.append(converter.to_omml(block.latex))


def add_table_block(document: Document, block: TableBlock) -> None:
    total_rows = len(block.rows) + (1 if block.headers else 0)
    total_cols = max(
        len(block.headers),
        max((len(row) for row in block.rows), default=0),
    )
    if total_rows == 0 or total_cols == 0:
        return

    table = document.add_table(rows=total_rows, cols=total_cols)
    table.style = "Table Grid"

    row_index = 0
    if block.headers:
        for column_index, value in enumerate(block.headers):
            cell = table.cell(0, column_index)
            cell.text = value
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if paragraph.runs:
                set_run_font(paragraph.runs[0], "宋体", 11, bold=True)
        row_index = 1

    for data_row in block.rows:
        for column_index, value in enumerate(data_row):
            cell = table.cell(row_index, column_index)
            cell.text = value
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if cell.paragraphs[0].runs:
                set_run_font(cell.paragraphs[0].runs[0], "宋体", 11)
        row_index += 1

    document.add_paragraph()


def add_image_block(document: Document, block: ImageBlock) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    width = Cm(block.width_cm) if block.width_cm else None
    run.add_picture(BytesIO(block.image_bytes), width=width)
    if block.caption.strip():
        caption = document.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_run = caption.add_run(block.caption)
        set_run_font(caption_run, font_name="宋体", font_size=10)


def add_numbered_equation(document: Document, converter: OmmlConverter, latex: str, number: str) -> None:
    table = document.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    remove_table_borders(table)

    widths = [Cm(1.5), Cm(11.7), Cm(2.2)]
    row = table.rows[0]
    for cell, width in zip(row.cells, widths, strict=True):
        cell.width = width
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.paragraphs[0].paragraph_format.space_before = Pt(0)
        cell.paragraphs[0].paragraph_format.space_after = Pt(0)

    center_paragraph = row.cells[1].paragraphs[0]
    center_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    center_run = center_paragraph.add_run()
    center_run._element.append(converter.to_omml(latex))

    number_paragraph = row.cells[2].paragraphs[0]
    number_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    number_run = number_paragraph.add_run(number)
    set_run_font(number_run, font_name="Times New Roman", font_size=12)

    document.add_paragraph()


def remove_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_borders = tbl_pr.first_child_found_in("w:tblBorders")
    if tbl_borders is None:
        tbl_borders = OxmlElement("w:tblBorders")
        tbl_pr.append(tbl_borders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge_element = tbl_borders.find(qn(f"w:{edge}"))
        if edge_element is None:
            edge_element = OxmlElement(f"w:{edge}")
            tbl_borders.append(edge_element)
        edge_element.set(qn("w:val"), "nil")
